"""
Build results_dem_915mhz.docx from the markdown report + PNG charts.
Run: python make_report_docx.py
"""
import re, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = '/home/user/Ray-Tracing---Sionna-RT'
MD   = os.path.join(BASE, 'results_dem_915mhz.md')
OUT  = os.path.join(BASE, 'results_dem_915mhz.docx')

# Chart captions and placement (after which section heading)
CHARTS = [
    ('dem_chart_overall_metrics.png',
     'Figure 1 — Overall accuracy: RMSE, Bias and R² per combining method (DEM terrain).',
     '5. Validation Against Measured Path Loss'),
    ('dem_chart_perband_rmse.png',
     'Figure 2 — Per-distance-band RMSE for Best ON, Incoh ON and FSPL reference.',
     '5.2 Per-Band RMSE'),
    ('dem_chart_ray_classification.png',
     'Figure 3 — Ray type breakdown: diffraction dominates at short range, multi-reflection at long range.',
     '6. Ray Propagation Analysis'),
    ('dem_chart_cumulative.png',
     'Figure 4 — Cumulative R² and RMSE vs distance threshold (Incoh ON, scattering ON).',
     '7. Cumulative Distance Evaluation'),
    ('dem_chart_calibration_compare.png',
     'Figure 5 — Scatter coefficient comparison: global S=0.70 vs per-material S values.',
     '8. Calibration'),
    ('dem_chart_cell8e.png',
     'Figure 6 — DEM + Roads (Run 3): cumulative RMSE and R² vs distance threshold.',
     '9. DEM + Roads Simulation'),
]

# ── helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def add_figure(doc, img_path, caption):
    if not os.path.exists(img_path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(6.0))
    add_caption(doc, caption)
    doc.add_paragraph()  # spacer

def parse_md_table(lines):
    """Return (headers, rows) from markdown table lines."""
    headers, rows = [], []
    for line in lines:
        line = line.strip()
        if not line or set(line.replace('|','').replace('-','').replace(':','').strip()) == set():
            continue  # separator row
        cells = [c.strip() for c in line.strip('|').split('|')]
        if not headers:
            headers = cells
        elif re.match(r'^[-:| ]+$', line):
            continue
        else:
            rows.append(cells)
    return headers, rows

def add_md_table(doc, headers, rows):
    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers[:n_cols]):
        cell = hdr.cells[i]
        cell.text = h.replace('**','').replace('`','')
        set_cell_bg(cell, '1F3864')
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'EBF0F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row[:n_cols]):
            cell = tr.cells[ci]
            txt  = val.replace('**','').replace('`','')
            cell.text = txt
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                if '**' in val:
                    run.bold = True

    doc.add_paragraph()  # spacer after table

# ── parse markdown ────────────────────────────────────────────────────────────
with open(MD, encoding='utf-8') as f:
    lines = f.readlines()

# ── build document ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# Default body font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title block (first 5 lines of markdown)
title_lines = [l.strip() for l in lines[:6] if l.strip()]
title_para = doc.add_heading(title_lines[0].lstrip('#').strip(), level=0)
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
for tl in title_lines[1:]:
    p = doc.add_paragraph(tl.lstrip('*').rstrip('*'))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
doc.add_paragraph()

# Track which charts have been inserted
charts_inserted = set()

def maybe_insert_chart(doc, heading_text):
    for img, caption, trigger in CHARTS:
        if img in charts_inserted:
            continue
        if trigger.lower() in heading_text.lower():
            add_figure(doc, os.path.join(BASE, img), caption)
            charts_inserted.add(img)

# Parse body
i = 6
table_buf = []
in_table  = False
in_code   = False
code_buf  = []

while i < len(lines):
    raw  = lines[i]
    line = raw.rstrip()
    i   += 1

    # Code block
    if line.strip().startswith('```'):
        if not in_code:
            in_code = True
            code_buf = []
        else:
            in_code = False
            p = doc.add_paragraph('\n'.join(code_buf))
            p.style = doc.styles['No Spacing']
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(8.5)
        continue
    if in_code:
        code_buf.append(line)
        continue

    # Table
    if line.startswith('|'):
        table_buf.append(line)
        continue
    if table_buf:
        headers, rows = parse_md_table(table_buf)
        if headers:
            add_md_table(doc, headers, rows)
        table_buf = []

    # Headings
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level = len(m.group(1))
        text  = m.group(2).strip()
        h = doc.add_heading(text, level=min(level, 4))
        maybe_insert_chart(doc, text)
        continue

    # HR
    if re.match(r'^---+$', line.strip()):
        doc.add_paragraph('─' * 60).runs[0].font.size = Pt(8)
        continue

    # Empty line
    if not line.strip():
        continue

    # Normal paragraph — strip markdown bold/italic/code markers
    text = line
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'`(.+?)`',       r'\1', text)
    text = re.sub(r'^[0-9]+\.\s+',  '',    text)  # numbered list
    text = re.sub(r'^[-*]\s+',       '',    text)  # bullet

    if text.strip():
        doc.add_paragraph(text.strip())

# Flush any remaining table
if table_buf:
    headers, rows = parse_md_table(table_buf)
    if headers:
        add_md_table(doc, headers, rows)

doc.save(OUT)
print(f'Saved: {OUT}')
